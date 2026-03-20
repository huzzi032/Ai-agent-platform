"""Document processing service for extracting text from various file types."""
import os
import io
import re
import requests
from typing import Optional, Dict, Any
from pathlib import Path
import logging

# PDF processing
try:
    from pypdf import PdfReader
except ImportError:
    from PyPDF2 import PdfReader

# Image processing
try:
    from PIL import Image
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process various document types and extract text."""
    
    SUPPORTED_TYPES = {
        'pdf': ['.pdf'],
        'image': ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'],
        'text': ['.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm'],
        'document': ['.docx', '.doc']
    }
    
    def __init__(self):
        """Initialize document processor."""
        self.supported_extensions = []
        for extensions in self.SUPPORTED_TYPES.values():
            self.supported_extensions.extend(extensions)
    
    def get_file_type(self, filename: str) -> Optional[str]:
        """Determine file type from extension."""
        ext = Path(filename).suffix.lower()
        for file_type, extensions in self.SUPPORTED_TYPES.items():
            if ext in extensions:
                return file_type
        return None
    
    def is_supported(self, filename: str) -> bool:
        """Check if file type is supported."""
        return self.get_file_type(filename) is not None
    
    def process_file(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """Process a file and extract text."""
        file_type = self.get_file_type(original_filename)
        
        if not file_type:
            raise ValueError(f"Unsupported file type: {original_filename}")
        
        try:
            if file_type == 'pdf':
                return self._process_pdf(file_path, original_filename)
            elif file_type == 'image':
                return self._process_image(file_path, original_filename)
            elif file_type == 'text':
                return self._process_text(file_path, original_filename)
            elif file_type == 'document':
                return self._process_document(file_path, original_filename)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error processing file {original_filename}: {e}")
            raise
    
    def process_url(self, url: str) -> Dict[str, Any]:
        """Process content from a URL."""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            content_type = response.headers.get('content-type', '')
            
            # Try to extract text based on content type
            if 'text/html' in content_type:
                text = self._extract_text_from_html(response.text)
            elif 'application/pdf' in content_type:
                # Save to temp file and process
                temp_path = f"/tmp/temp_pdf_{os.urandom(4).hex()}.pdf"
                with open(temp_path, 'wb') as f:
                    f.write(response.content)
                result = self._process_pdf(temp_path, "downloaded.pdf")
                os.remove(temp_path)
                return result
            else:
                text = response.text
            
            return {
                "text": text,
                "filename": url.split('/')[-1] or "web_content",
                "file_type": "url",
                "metadata": {
                    "source_url": url,
                    "content_type": content_type,
                    "content_length": len(response.content)
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing URL {url}: {e}")
            raise
    
    def process_text_input(self, text: str, title: str = "user_input") -> Dict[str, Any]:
        """Process raw text input."""
        return {
            "text": text,
            "filename": title,
            "file_type": "text",
            "metadata": {
                "char_count": len(text),
                "word_count": len(text.split())
            }
        }
    
    def _process_pdf(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """Extract text from PDF."""
        try:
            reader = PdfReader(file_path)
            text_parts = []
            metadata = {
                "total_pages": len(reader.pages),
                "extracted_pages": 0
            }
            
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {i+1} ---\n{page_text}")
                        metadata["extracted_pages"] += 1
                except Exception as e:
                    logger.warning(f"Error extracting page {i+1}: {e}")
            
            full_text = "\n\n".join(text_parts)
            
            # Extract PDF metadata if available
            if reader.metadata:
                metadata["pdf_info"] = {
                    k: str(v) for k, v in reader.metadata.items() if v
                }
            
            return {
                "text": full_text,
                "filename": original_filename,
                "file_type": "pdf",
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise
    
    def _process_image(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """Extract text from image using OCR."""
        if not TESSERACT_AVAILABLE:
            raise ImportError("OCR not available. Install pytesseract and tesseract-ocr.")
        
        try:
            image = Image.open(file_path)
            
            # Get image metadata
            metadata = {
                "format": image.format,
                "mode": image.mode,
                "size": image.size,
                "width": image.width,
                "height": image.height
            }
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            return {
                "text": text,
                "filename": original_filename,
                "file_type": "image",
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing image: {e}")
            raise
    
    def _process_text(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """Process text files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            text = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("Could not decode file with any encoding")
            
            # Get file stats
            stat = os.stat(file_path)
            
            return {
                "text": text,
                "filename": original_filename,
                "file_type": "text",
                "metadata": {
                    "encoding": encoding,
                    "char_count": len(text),
                    "word_count": len(text.split()),
                    "line_count": len(text.splitlines()),
                    "file_size": stat.st_size
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing text file: {e}")
            raise
    
    def _process_document(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """Process Word documents."""
        try:
            # Try to use python-docx if available
            try:
                from docx import Document
                doc = Document(file_path)
                
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        text_parts.append(" | ".join(row_text))
                
                full_text = "\n\n".join(text_parts)
                
                return {
                    "text": full_text,
                    "filename": original_filename,
                    "file_type": "document",
                    "metadata": {
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables),
                        "char_count": len(full_text)
                    }
                }
                
            except ImportError:
                # Fallback: try to extract as text
                logger.warning("python-docx not available, trying text extraction")
                return self._process_text(file_path, original_filename)
                
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            raise
    
    def _extract_text_from_html(self, html: str) -> str:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except ImportError:
            # Fallback: simple regex
            text = re.sub('<[^<]+?>', '', html)
            return text


# Singleton instance
_processor = None


def get_document_processor() -> DocumentProcessor:
    """Get or create document processor singleton."""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
